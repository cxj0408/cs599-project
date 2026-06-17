from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "智能旅行助手课程报告_最终版.docx"


def east_asia(run, font="Microsoft YaHei", size=10.5):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(7)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    east_asia(r)
    return p


def no_indent_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    east_asia(r)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    east_asia(r, "Consolas", 8.5)
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    east_asia(r, size=9.3)
    r.bold = bold


def table(doc, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, text in enumerate(rows[0]):
        set_cell(t.rows[0].cells[j], text, True)
        shade(t.rows[0].cells[j], "E8EEF5")
    for row in rows[1:]:
        cells = t.add_row().cells
        for j, text in enumerate(row):
            set_cell(cells[j], text)
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    east_asia(r, size=9)
    r.italic = True
    r.font.color.rgb = RGBColor(85, 85, 85)


def title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("智能旅行助手课程设计报告")
    east_asia(r, size=24)
    r.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 Multi-Agent 与高德地图 MCP 的 AI Travel Agent")
    east_asia(r, size=12)
    r.font.color.rgb = RGBColor(85, 85, 85)

    table(doc, [
        ["项目项", "内容"],
        ["项目名称", "智能旅行助手（AI Travel Agent）"],
        ["项目方向", "Agentic AI 原生开发"],
        ["核心技术", "LangChain、LangGraph、MCP、DashScope 通义千问、高德地图 MCP、Streamlit"],
        ["交付形态", "CLI 命令行、Streamlit Web 界面、Markdown 下载"],
        ["文档说明", "本报告根据项目目录中的 README、源码、测试脚本、运行日志和示例截图整理，并按课程评分章节展开。"],
    ], widths=[3.1, 12.6])
    doc.add_page_break()


def main():
    doc = Document()
    configure_doc(doc)
    title_page(doc)

    heading(doc, "一、选题背景与设计思想（20分）", 1)
    heading(doc, "1.1 问题定义", 2)
    para(doc, "旅行规划看起来是一个普通的信息检索问题，但真正做起来会发现它其实是一个连续决策问题。用户不会只问“长沙有什么景点”，而是会把时间、预算、住宿位置、交通方式、兴趣偏好、同行人情况等条件放在一起，希望系统直接给出一套能落地执行的安排。例如项目中的示例需求是“长沙 3 日游，喜欢自然风光和历史文化，中等预算，住五一广场”，这句话里同时包含目的地、行程长度、兴趣约束、预算约束和住宿区域约束。系统如果只返回景点清单，实际上并没有解决用户的问题；它还需要考虑每天去哪里、景点之间如何移动、天气是否影响户外活动、酒店是否方便、餐饮和交通大概花多少钱。")
    para(doc, "因此，本项目把“旅行规划”定义为一个由大模型负责理解与编排、由外部工具负责事实查询的 Agent 任务。用户输入自然语言后，系统需要自动拆解任务，分别查询天气、搜索景点、推荐酒店、规划路线，并把这些结果重新组织成按天排列的行程。最终输出不仅要有文字描述，还要有结构化 JSON，便于 Web 页面渲染、预算汇总和 Markdown 下载。")

    heading(doc, "1.2 现有方案不足", 2)
    para(doc, "现有的旅行规划方式最大的问题不是“查不到信息”，而是信息太分散、太难整合。用户规划一次旅行时，通常要先去天气软件看目的地天气，再去地图软件查景点位置和路线，又要到酒店平台比较住宿位置和价格，还会翻攻略社区了解哪些景点值得去、哪些餐厅适合顺路安排。每个平台都能解决一小块问题，但没有一个统一视图能把这些信息自然地串成一条完整行程。用户最后往往是在多个网页和应用之间来回切换，手动复制地址、估算距离、调整顺序，规划过程很容易变成一项繁琐的整理工作。")
    para(doc, "如果直接使用普通大模型来生成行程，体验会轻松一些，但它又会带来另一个问题：模型本身并不真正知道实时天气、最新 POI 状态、具体地址和路线耗时。它可能会给出听起来很合理的建议，却把已经变化的信息当成事实，甚至编造并不存在的景点细节。对于旅行这种强依赖现实位置和时间的信息任务来说，这类错误会直接影响用户决策。传统攻略内容也存在类似局限，它们大多是静态文章，适合参考，但很难根据用户的具体日期、预算、住宿区域和兴趣点即时重组。也就是说，现有方案要么让用户自己承担大量整合成本，要么给出不够实时、不够可靠的自动化建议，距离“一站式、可执行、个性化”的智能旅行规划还有明显差距。")

    heading(doc, "1.3 项目价值", 2)
    para(doc, "本项目的价值在于把大模型从“给建议的人”变成“会调用工具完成任务的规划助手”。在系统中，大模型不再单纯依赖已有知识回答，而是通过 MCP 协议接入高德地图工具，实时查询天气、POI 和路线信息。这样做可以减少凭空生成的风险，也能让最终行程更贴近真实场景。用户看到的不只是几段泛泛的旅游建议，而是一份包含天气、酒店、景点、餐饮、交通和预算的完整计划。")
    para(doc, "从工程角度看，这个项目也比较适合作为课程设计，因为它不是单一页面或单一接口的堆叠，而是包含了 Agent 架构、工具调用、流式输出、配置管理、异常诊断、缓存设计和前端展示的完整链路。它能体现 Agentic AI 原生开发的特点：模型负责理解和决策，工具负责获取外部事实，代码负责把两者组织成稳定产品。")

    heading(doc, "1.4 设计思想与技术路线", 2)
    para(doc, "在设计思想上，本项目没有把所有能力都写进一个大而全的提示词里，而是把旅行规划拆成“总控编排、领域处理、工具调用、结果渲染”几个相对独立的层次。TripPlanner 承担门面模式的角色，对外只暴露 invoke() 和 stream() 两个简单接口，调用者不需要知道内部还有酒店、景点、天气等多个子 Agent，也不需要关心 MCP 工具如何加载。这样的设计让系统入口保持简单，同时把复杂性留在内部可控的位置。")
    para(doc, "底层连接部分使用单例模式管理 McpClientManager，通过 __new__ 和 _initialized 做双重保护，保证整个程序只维护一套 MCP 客户端和工具缓存。这个选择和 Streamlit 的运行机制有关：Web 页面每次交互都会触发脚本重新执行，如果不做单例和缓存，就可能反复创建连接、重复拉取工具列表，既浪费时间，也增加外部服务不稳定的概率。Config.create_llm() 则更接近工厂方法，把模型名称、API Key、温度、streaming 等构造细节集中封装，后续如果更换模型或调整参数，不需要在业务代码里到处修改。")
    para(doc, "Agent 之间的协作也体现了装饰器模式和策略思想。项目用 @tool 把 SpecialistAgent 包装成 LangChain Tool，使 HotelAgent、AttractionAgent、WeatherAgent 能够被总控 Agent 像普通工具一样调用；同时，config.py 中的 tool_domains 字典把 MCP 工具按 poi、weather、route 进行分组，运行时根据任务领域选择不同工具集合。这样做比让每个 Agent 持有全部工具更稳妥，既降低了误调用概率，也让每个子 Agent 的职责更清晰。SpecialistAgent.build() 则提供了类似模板方法的扩展点，目前三个专家 Agent 共享同一套构建流程，后续如果新增餐饮、交通或评价 Agent，也可以复用这个结构，只替换提示词和工具集合。")
    para(doc, "结果展示部分可以看作适配器思想的体现。Planner 输出的是统一 JSON，但用户最终看到的可能是 CLI 里的格式化文本，也可能是 Streamlit 页面里的天气卡片、每日行程 Tab 和预算指标。render.py 的作用就是把同一份结构化数据适配成不同展示形式。整体来看，项目采用这些设计模式不是为了堆概念，而是为了解决实际工程问题：连接要复用，模型创建要统一，Agent 职责要拆开，工具权限要受控，输出结果要能被不同界面稳定消费。")
    para(doc, "技术路线上，系统使用通义千问 qwen3-max 作为语言模型，使用 LangChain 和 LangGraph 构建 ReAct Agent，使用 langchain-mcp-adapters 连接阿里百炼高德地图 MCP 服务，使用 Streamlit 构建 Web 界面。用户侧既可以通过 Agent.py 在命令行演示，也可以通过 app.py 使用图形界面。输出侧则通过 prompts.py 约束 JSON Schema，再由 render.py 和 Web 组件负责展示。")

    heading(doc, "二、Specs 规格文档（20分）", 1)
    heading(doc, "2.1 Product Spec", 2)
    para(doc, "产品目标是做一个轻量但完整的智能旅行规划助手。它面向的用户不是专业旅行社，而是希望快速安排个人或小团队行程的普通用户。用户只需要输入目的地、日期范围、交通偏好、住宿偏好、旅行兴趣和补充要求，系统就要生成一份可以直接参考的旅行计划。这个计划应当包括每天的主要景点、推荐酒店、三餐建议、天气情况、路线交通方式和费用估算。")
    table(doc, [
        ["规格项", "具体内容"],
        ["目标用户", "希望快速生成国内城市旅行计划的普通游客、学生、小团队出行者。"],
        ["核心输入", "目的地城市、开始日期、结束日期、交通方式、住宿类型、旅行偏好、额外要求。"],
        ["核心输出", "按天组织的行程、天气信息、酒店信息、景点信息、餐饮建议、预算汇总和总体建议。"],
        ["主要交互", "Web 端通过侧边栏填写参数并点击开始规划；CLI 端通过预设输入演示流式输出。"],
        ["非功能要求", "输出结构稳定、工具调用过程可见、异常信息可诊断、结果可下载。"],
    ], widths=[3.2, 12.5])
    para(doc, "产品设计中比较重要的一点是，Web 页面不是简单显示模型原文，而是把结果拆成更适合阅读的模块。天气用卡片展示，每日行程用 Tab 切换，预算用指标卡汇总，原始结果可以保存为 Markdown。这让系统更像一个可用的小产品，而不是一次性的聊天记录。")

    heading(doc, "2.2 Architecture Spec", 2)
    para(doc, "系统架构分为四层：用户界面层、Agent 编排层、MCP 工具层和外部服务层。用户界面层负责接收输入和展示结果；Agent 编排层负责理解需求、拆解任务并调用工具；MCP 工具层负责把地图能力包装成 LangChain Tool；外部服务层则是阿里百炼托管的高德地图 MCP 服务。")
    code(doc, """用户界面层
├─ Agent.py：命令行演示入口
└─ app.py：Streamlit Web 入口
        │
        ▼
TripPlanner 总控 Agent
├─ search_hotel       → HotelAgent       → MCP: maps_text_search / maps_search_detail
├─ search_attraction  → AttractionAgent  → MCP: maps_text_search / maps_search_detail
├─ query_weather      → WeatherAgent     → MCP: maps_weather
└─ route tools        → MCP: maps_direction_*_by_address
        │
        ▼
McpClientManager 单例
├─ MultiServerMCPClient
├─ 按 poi / weather / route 分发工具
└─ RedisMcpCache 可选缓存
        │
        ▼
阿里百炼高德地图 MCP 服务""")
    para(doc, "TripPlanner 并不直接写死每一种底层 API 调用，而是通过 McpClientManager 获取工具列表，再把其中属于 poi、weather、route 的工具分配给不同模块。这让系统在结构上具有一定弹性：如果 MCP 服务新增工具，只需要在配置中更新领域映射，Agent 层就可以按需使用。")

    heading(doc, "2.3 API Spec", 2)
    para(doc, "项目内部接口主要围绕“规划器、专家 Agent、MCP 客户端和渲染器”展开。TripPlanner 是对外的核心调用入口；SpecialistAgent 是领域能力的统一封装；McpClientManager 负责连接外部工具；render.py 负责把模型输出转成可展示数据。")
    table(doc, [
        ["接口/函数", "输入", "输出", "说明"],
        ["TripPlanner.invoke(user_input)", "自然语言旅行需求", "完整结果文本", "非流式调用，适合测试或一次性生成。"],
        ["TripPlanner.stream(user_input)", "自然语言旅行需求", "AsyncIterator[str]", "流式返回 token 和工具调用状态，适合 CLI/Web 展示。"],
        ["SpecialistAgent.invoke(user_input)", "领域查询语句", "领域查询结果", "封装单领域 ReAct Agent。"],
        ["McpClientManager.get_tools_for(domain)", "poi/weather/route", "BaseTool 列表", "按领域过滤 MCP 工具，减少误调用。"],
        ["render.parse_plan(text)", "混合文本", "dict | None", "从模型输出中截取并解析 JSON。"],
        ["app.build_prompt(...)", "表单参数", "自然语言 prompt", "把 Web 结构化输入转为 Agent 能理解的描述。"],
    ], widths=[4.1, 3.2, 3.2, 5.2])

    heading(doc, "2.4 输出规格", 2)
    para(doc, "输出规格采用 JSON Schema 思路进行约束。这样做的原因是，旅行计划最终要被 Web 页面拆分渲染，如果只是普通自然语言，程序很难稳定识别哪一段是天气、哪一段是酒店、哪一段是预算。通过在 PLANNER_AGENT_PROMPT 中明确字段，系统可以把大模型的生成结果转换为更稳定的数据结构。")
    code(doc, """{
  "city": "城市名称",
  "start_date": "YYYY-MM-DD",
  "end_date": "YYYY-MM-DD",
  "days": [{
    "date": "YYYY-MM-DD",
    "hotel": {"name": "...", "address": "...", "estimated_cost": 400},
    "attractions": [{"name": "...", "visit_duration": 120, "ticket_price": 60}],
    "meals": [{"type": "breakfast/lunch/dinner", "estimated_cost": 30}]
  }],
  "weather_info": [{"date": "...", "day_weather": "晴", "day_temp": 25}],
  "budget": {"total_attractions": 180, "total_hotels": 1200, "total": 2060}
}""")

    heading(doc, "三、系统架构与设计（15分）", 1)
    heading(doc, "3.1 核心架构图", 2)
    para(doc, "系统核心架构可以理解为一条从用户意图到外部工具、再到结构化结果的链路。用户输入不是直接交给工具，而是先交给总控 Agent，由它判断需要调用哪些子 Agent 和路线工具。子 Agent 得到任务后再调用受限的 MCP 工具，最后由 Planner 汇总所有结果。")
    code(doc, """┌──────────────────────────────┐
│ 用户输入：城市 / 日期 / 偏好 / 预算 │
└──────────────┬───────────────┘
               ▼
┌──────────────────────────────┐
│ app.py / Agent.py 构造 Prompt │
└──────────────┬───────────────┘
               ▼
┌──────────────────────────────┐
│ TripPlanner 总控 Agent        │
│ 决定调用天气、酒店、景点、路线工具 │
└─────┬────────┬────────┬──────┘
      ▼        ▼        ▼
 Weather   Hotel    Attraction      Route MCP Tools
 Agent     Agent    Agent           walking/driving/transit
      └────────┬────────┘              │
               ▼                       ▼
        McpClientManager 单例 + 缓存包装
               ▼
        阿里百炼高德地图 MCP 服务
               ▼
        JSON 旅行计划 → Web/CLI 渲染""")

    heading(doc, "3.2 Agent 交互流程", 2)
    para(doc, "一次完整交互从 Web 表单开始。用户填写目的地、时间范围和偏好后，app.py 的 build_prompt 会把这些表单字段组织成一句更自然的旅行需求。随后 TripPlanner.stream 接收这段需求并启动 LangGraph Agent。由于系统使用流式事件，页面可以在 Agent 工作时看到“查询天气中”“搜索酒店中”“搜索景点中”“规划路线中”等状态，而不是让用户盯着空白页面等待。")
    para(doc, "在内部流程上，TripPlanner 首次运行会调用 build 方法。这个方法先通过 McpClientManager 获取三类工具：poi 工具用于景点和酒店搜索，weather 工具用于天气查询，route 工具用于路线规划。然后系统创建三个 SpecialistAgent，并把它们包装成 search_hotel、search_attraction、query_weather 三个 LangChain Tool。这样总控 Agent 看到的是统一工具接口，内部却可以由子 Agent 完成更细的领域推理。")
    para(doc, "这种设计避免了单个 Agent 既要理解全局行程、又要掌握所有工具细节的问题。总控 Agent 只关注“什么时候调用哪个能力”，子 Agent 只关注“如何完成单一领域查询”。从工程维护角度看，它比一个超长 Prompt 更容易调试，也更容易扩展。")

    heading(doc, "3.3 数据流设计", 2)
    para(doc, "数据流设计的重点是把“自由文本输入”和“结构化页面输出”连接起来。用户输入最初是表单字段，进入 Agent 后变成自然语言需求；Agent 输出最初是流式文本，最后又被 parse_plan 提取为 JSON；Web 页面再根据 JSON 的不同字段渲染成天气卡片、每日行程、预算指标和下载内容。")
    code(doc, """Web 表单数据
  → build_prompt()
  → TripPlanner.stream()
  → astream_events(v2)
  → on_tool_start：显示工具状态
  → on_chat_model_stream：收集模型输出
  → parse_plan() 提取 JSON
  → st.session_state 保存 plan_data / plan_raw
  → Streamlit 页面渲染 + Markdown 下载""")
    para(doc, "这里的 st.session_state 很关键，因为 Streamlit 每次交互都会重新运行脚本。如果不保存 plan_data 和 plan_raw，用户点击按钮后页面刷新，结果就可能丢失。项目使用 session_state 持久化结果，再通过 st.rerun 触发展示逻辑，使生成过程和结果展示分开。")

    heading(doc, "3.4 Demo 截图说明", 2)
    para(doc, "项目目录中提供了三张示例截图，展示了系统的 Web 形态和生成效果。报告中保留这些截图，是为了让架构设计不只停留在文字层面，也能体现最终用户看到的界面。")
    for name in ["示例1.png", "示例2.png", "示例3.png"]:
        img = ROOT / name
        if img.exists():
            doc.add_picture(str(img), width=Cm(13.5))
            caption(doc, f"项目运行示例：{name}")

    heading(doc, "四、关键实现与代码展示（15分）", 1)
    heading(doc, "4.1 Agent 核心循环", 2)
    para(doc, "Agent 核心循环位于 TripPlanner.stream。这个函数不是简单等待最终答案，而是监听 LangGraph 的事件流。模型生成普通 token 时，系统把 token 继续传给前端；工具开始调用时，系统把工具名映射成中文状态提示；工具结束时则保持安静，避免页面被大量中间信息刷屏。")
    code(doc, """async for event in self._agent.astream_events(
    {"messages": [{"role": "user", "content": user_input}]},
    version="v2",
):
    kind = event.get("event", "")
    if kind == "on_chat_model_stream":
        content = event["data"]["chunk"].content
        content = _TOOL_CALL_PATTERN.sub("", content)
        if content.strip():
            yield content
    elif kind == "on_tool_start":
        name = event.get("name", "unknown")
        emoji, label = TOOL_LABELS.get(name, ("tool", name))
        yield f"\\n{emoji} {label}...\\n" """)
    para(doc, "这段代码体现了 Agent 产品化时一个很实际的问题：用户不仅关心最终结果，也关心系统是不是正在工作。流式状态可以降低等待焦虑，同时也便于调试工具调用顺序。")

    heading(doc, "4.2 子 Agent 包装为工具", 2)
    para(doc, "本项目比较有代表性的实现，是把子 Agent 包装成 Tool。HotelAgent、AttractionAgent 和 WeatherAgent 内部本身也是 Agent，但对 TripPlanner 来说，它们只是三个可调用函数。这种“Agent as Tool”的写法让系统层次更清楚，也使总控 Agent 的提示词更容易管理。")
    code(doc, """@tool
async def search_hotel(query: str) -> str:
    return await self._hotel_agent.invoke(query)

@tool
async def search_attraction(query: str) -> str:
    return await self._attraction_agent.invoke(query)

@tool
async def query_weather(query: str) -> str:
    return await self._weather_agent.invoke(query)

all_tools = [search_hotel, search_attraction, query_weather, *route_tools]""")
    para(doc, "这样做还有一个好处：权限边界更明确。天气 Agent 只拿天气工具，景点和酒店 Agent 只拿 POI 工具，路线工具则由 Planner 直接持有。即使模型在某个子任务中产生偏差，也不容易调用到不相关工具。")

    heading(doc, "4.3 MCP 客户端与缓存", 2)
    para(doc, "mcp_client.py 中的 McpClientManager 使用单例模式。这样设计主要是为了适配 Streamlit 的运行方式：Streamlit 每次用户交互都会重新执行脚本，如果每次都重新创建 MCP 客户端和工具列表，不仅耗时，也容易造成重复连接。单例和工具缓存让系统只在首次使用时加载工具，之后复用已有结果。")
    code(doc, """class McpClientManager:
    _instance: "McpClientManager | None" = None

    def __new__(cls) -> "McpClientManager":
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance""")
    para(doc, "项目还加入了 RedisMcpCache，用于缓存天气和 POI 详情等相对稳定、重复率较高的 MCP 调用结果。缓存采用 fail-open 设计，也就是 Redis 不可用时只关闭缓存，不影响核心规划功能。运行日志中出现过“Redis MCP cache disabled: No module named 'redis'”，但系统仍然可以继续走 MCP 调用链，这说明缓存层没有绑死核心流程。")

    heading(doc, "4.4 配置文件与上游兼容修复", 2)
    para(doc, "config.py 不只是存放 API Key，它还统一管理模型名称、温度、MCP URL、传输协议、缓存开关和工具领域映射。项目中比较典型的工程处理是对 ChatTongyi 的流式 tool_calls 差分逻辑做了 Monkey-Patch。原因是上游库在处理流式工具调用增量时，可能直接访问 prev_function['name'] 或 prev_function['arguments']，而某些 chunk 并不包含这些字段，最终触发 KeyError。")
    code(doc, """if "name" in function and "name" in prev_function:
    function["name"] = function["name"].replace(prev_function["name"], "")
if "arguments" in function and "arguments" in prev_function:
    function["arguments"] = function["arguments"].replace(
        prev_function["arguments"], ""
    )""")
    para(doc, "这个修复虽然代码不长，但体现了 AI 应用开发中经常遇到的问题：很多故障并不在业务逻辑本身，而在模型 SDK、协议适配器、流式事件或外部服务兼容性上。能把这些问题定位并修掉，是项目能真正运行起来的重要部分。")

    heading(doc, "4.5 Web 端实现", 2)
    para(doc, "Web 端使用 Streamlit 实现。侧边栏负责收集旅行参数，主区域负责展示生成结果。为了避免每次点击都重新创建规划器，app.py 使用 @st.cache_resource 缓存 TripPlanner 实例；为了保留生成结果，使用 st.session_state 保存 plan_data、plan_raw 和 status_lines。")
    para(doc, "页面展示上，天气信息被渲染为卡片，每日行程放入 Tab，预算用 st.metric 呈现，用户还可以下载 Markdown。这样的设计虽然不复杂，但已经具备一个小型 AI 工具产品的基本体验：输入明确、状态可见、输出结构清晰、结果可以保存。")

    heading(doc, "五、测试与评估（10分）", 1)
    heading(doc, "5.1 功能测试", 2)
    para(doc, "项目目录中包含多种测试和诊断脚本，说明开发过程并不是只跑一次 Demo，而是围绕环境、协议、MCP 服务和配置做了多轮验证。check_env.py 用于检查 Python 版本、关键依赖、API Key 和网络连通性；test_mcp.py 用于验证 McpClientManager 能否成功获取高德地图工具；test_adapter_usage.py 则比较了不同 MCP 连接方式。")
    table(doc, [
        ["测试项", "脚本/位置", "验证目标"],
        ["环境诊断", "check_env.py", "检查 Python、依赖、API Key、网络与 MCP 端点。"],
        ["MCP 连接", "test_mcp.py", "验证是否能获取 MCP 工具列表。"],
        ["传输协议", "test_adapter_usage.py / test_transport.py", "比较 streamable-http、SSE、direct MCP 等连接方式。"],
        ["配置验证", "verify_config.py", "检查 MCP_TRANSPORT、MCP_URL、缓存开关等配置。"],
        ["Web 功能", "app.py + 示例截图", "验证表单输入、状态展示、结果渲染和下载。"],
    ], widths=[3, 5.1, 7.6])

    heading(doc, "5.2 Agent 行为评估", 2)
    para(doc, "Agent 的评估不能只看回答有没有文字，还要看它是否按预期调用工具、是否生成可解析 JSON、是否覆盖旅行计划所需字段。对于本项目来说，天气、酒店、景点这类事实型信息必须走工具查询，不能让模型凭空回答；最终 JSON 必须能被 parse_plan 成功解析，否则 Web 页面就无法稳定展示。")
    para(doc, "另外，行程本身也要评估合理性。每天是否安排 2 到 3 个景点，是否包含三餐，住宿是否和活动区域匹配，预算是否有分项和总计，天气信息是否覆盖每一天，这些都可以作为行为评估指标。相比普通问答任务，旅行规划更强调“可执行性”，所以评估时不能只看语言是否通顺，还要看安排是否能真正拿来使用。")

    heading(doc, "5.3 Benchmark 设计", 2)
    para(doc, "为了后续持续改进，可以建立一组标准测试输入。例如选择长沙、杭州、成都、北京等不同城市，分别设置 1 日游、3 日游、5 日游，不同预算和不同偏好。每次修改提示词、模型或 MCP 连接方式后，都用同一组输入测试，观察 JSON 可解析率、工具调用覆盖率、字段完整率和总耗时。")
    table(doc, [
        ["评估维度", "指标", "含义"],
        ["结构稳定性", "JSON 可解析率", "模型输出能否被 render.parse_plan 正常解析。"],
        ["工具使用", "事实型任务工具调用覆盖率", "天气、酒店、景点是否实际调用 MCP 工具。"],
        ["内容完整性", "字段覆盖率", "days、weather_info、hotel、attractions、meals、budget 是否齐全。"],
        ["行程合理性", "路线与时间冲突数", "景点距离是否过远，游玩时间是否明显不合理。"],
        ["性能表现", "首 token 时间 / 总耗时", "用户等待体验是否可接受。"],
    ], widths=[3, 4.4, 8.3])

    heading(doc, "5.4 当前测试观察", 2)
    para(doc, "从现有日志看，Streamlit 服务曾在 127.0.0.1:8501 成功启动，说明 Web 层可以运行。日志中也记录了 Redis 未安装导致缓存关闭的问题，但因为缓存设计为 fail-open，所以这不是致命故障。另一个更关键的问题是 MCP 服务曾返回 500 Internal Server Error，这通常需要检查阿里百炼控制台是否已经开通高德地图 MCP 服务、API Key 是否有效、传输协议是否和当前依赖版本匹配。项目中保留了多份排查文档和测试脚本，说明这部分已经纳入工程诊断范围。")

    heading(doc, "六、系统升级与扩展（10分）", 1)
    heading(doc, "6.1 可扩展架构", 2)
    para(doc, "当前系统的扩展点比较清晰。新增能力时，可以新增一个 SpecialistAgent，也可以直接把新的 MCP 工具加入 TripPlanner 的工具列表。例如要增加餐饮推荐，可以创建 FoodAgent，并给它分配餐厅搜索、周边搜索和详情查询工具；要增加火车票或机票能力，可以创建 TransportAgent；要增加行程复盘和质量检查，可以创建 EvaluatorAgent。")
    para(doc, "这种架构比把所有逻辑写在一个函数里更适合长期迭代。每个 Agent 都有独立的提示词和工具集合，调试时可以单独观察某个领域的输出。McpClientManager 负责统一连接和缓存，也避免了每个模块都重复写连接逻辑。")

    heading(doc, "6.2 下一阶段计划", 2)
    para(doc, "下一阶段首先要解决稳定性问题。MCP 连接涉及外部服务、API Key、传输协议和依赖版本，任何一环不稳定都会影响 Demo，所以需要加入更明确的超时、重试和错误提示。其次要加强质量评估，建立标准测试集，避免每次只凭一次成功结果判断系统好坏。")
    para(doc, "产品层面可以继续增强交互能力，例如允许用户对某一天行程进行局部调整，而不是重新生成全部计划；允许用户锁定某个酒店或景点，让 Agent 只重排剩余部分；允许导出 PDF、日历或地图链接。部署层面则可以补充 Dockerfile、环境变量模板和云端部署说明，使项目更容易复现。")

    heading(doc, "6.3 AI 能力演进路径", 2)
    para(doc, "从能力演进看，本项目目前处于“工具增强型 Agent”阶段，核心目标是让模型能够查询真实工具并生成结构化结果。下一步可以进入“可评估 Agent”阶段，在生成后增加自检：比如检查每天景点数量是否合理、雨天是否安排过多户外项目、预算是否超过用户要求。再往后可以加入记忆模块，让系统记住用户喜欢自然风光、讨厌赶路、偏好市中心住宿等长期偏好。")
    para(doc, "更长期的方向是形成自治式旅行助理。它不仅能规划，还能接入预订、退改、实时天气提醒、交通延误提醒和临时改线。那时 Agent 的角色就不只是生成一份计划，而是陪伴用户完成整个出行过程。")

    heading(doc, "七、课程总结（10分）", 1)
    heading(doc, "7.1 个人收获", 2)
    para(doc, "通过这个项目，我对 AI 应用开发的理解发生了比较明显的变化。最开始容易把大模型应用理解成“写好 Prompt，让模型回答”，但真正做旅行规划时会发现，Prompt 只是入口，系统能不能用，更多取决于工具接入、数据结构、异常处理和界面展示。一个看似简单的“帮我规划三日游”，背后其实包含任务拆解、事实查询、路线组织、预算计算和结果渲染。")
    para(doc, "项目中最有收获的部分是 Multi-Agent 的拆分。总控 Agent 不需要亲自处理所有细节，子 Agent 也不需要理解全局目标。它们通过工具接口连接起来，各自承担一部分职责。这种设计让我更直观地理解了 Agentic AI 和传统问答系统的区别：Agent 不是只生成文本，而是在一组工具和约束中完成任务。")

    heading(doc, "7.2 工程思维转变", 2)
    para(doc, "这个项目也让我意识到，AI 项目不是“模型效果好就结束”。实际开发中会遇到 API Key 配置、MCP 服务开通、依赖版本兼容、Streamlit rerun、Redis 缓存缺失、流式 tool_calls bug 等一系列工程问题。很多时候，真正花时间的不是写生成逻辑，而是把链路打通、把错误定位清楚、把结果稳定展示出来。")
    para(doc, "因此，我对工程化的理解从“实现功能”进一步变成了“让功能可运行、可诊断、可扩展”。报告中的测试脚本、配置管理和异常提示虽然不像界面那样直观，但它们决定了项目能不能在不同环境中复现。")

    heading(doc, "7.3 对课程的建议", 2)
    para(doc, "如果后续课程继续围绕 Agentic AI 展开，建议增加一些工具调用和 MCP 的实操内容。相比单纯讲 Prompt，工具调用更能体现大模型应用和真实系统之间的关系。也建议课程中加入标准化评估环节，例如要求每个项目提供一组固定测试输入，说明成功率、失败样例和改进方向。这样可以避免只展示一次成功 Demo，而忽略系统稳定性。")
    para(doc, "另外，AI IDE 的使用也很适合加入课程过程。学生可以用 AI IDE 辅助阅读源码、生成测试、定位日志和整理文档，但同时也要学会判断 AI 生成内容是否可靠。这样的训练更接近真实开发场景：AI 可以提高效率，但最终仍然需要开发者理解系统结构并做出工程判断。")

    heading(doc, "附录：项目文件对应关系", 1)
    table(doc, [
        ["文件", "作用"],
        ["Agent.py", "CLI 入口，演示流式/非流式生成旅行计划。"],
        ["app.py", "Streamlit Web UI，负责表单输入、流式收集、结果展示和 Markdown 下载。"],
        ["agents/planner.py", "TripPlanner 总控 Agent，编排子 Agent 和路线 MCP 工具。"],
        ["agents/specialist.py", "领域专家 Agent 封装，支持天气、景点、酒店等单职责任务。"],
        ["mcp_client.py", "MCP 客户端单例，加载和分发高德地图 MCP 工具。"],
        ["mcp_cache.py", "Redis 缓存封装，缓存天气和 POI 详情等高频工具结果。"],
        ["config.py", "全局配置与 ChatTongyi Monkey-Patch。"],
        ["prompts.py", "Planner、Weather、Attraction、Hotel 的系统提示词。"],
        ["render.py", "JSON 解析与 CLI 格式化渲染。"],
        ["test_*.py / check_*.py", "连接、协议、环境和兼容性测试脚本。"],
    ], widths=[4.5, 11.2])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
